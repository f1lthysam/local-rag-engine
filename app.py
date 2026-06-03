"""
app.py — Flask server with Admin Dashboard + RAG Studio

Routes:
    /               → Dashboard (tenant management)
    /studio         → RAG Studio (existing test UI)
    /ingest-url     → Scrape & index URL
    /upload-file    → Upload & index file
    /clear-chat     → Clear chat history
    /select-dataset → Switch dataset filter
    /api/tenants    → CRUD for tenants (JSON API)
"""

import warnings
warnings.filterwarnings("ignore")
import os, json, logging, secrets
from uuid import uuid4
from urllib.parse import urlparse
from pathlib import Path
os.environ["PYTHONWARNINGS"] = "ignore"
logging.disable(logging.CRITICAL)

from flask import Flask, jsonify, redirect, render_template, request, send_from_directory, session, url_for
from werkzeug.utils import secure_filename
from pypdf import PdfReader
from populate_database import add_to_chroma, load_documents, split_documents
import query_data
from scrape_web import scrape_and_save, scrape_full_website

app = Flask(__name__, template_folder="templates")
app.config["TEMPLATES_AUTO_RELOAD"] = True
app.secret_key = os.getenv("FLASK_SECRET_KEY", "local-rag-dev-secret")

MODEL_NAME  = "gemini-3.1-flash-lite"
MAX_CHAT_TURNS = 20
CHAT_SESSIONS  = {}
DATA_PATH      = "data"
TENANTS_FILE   = Path("tenants.json")
HISTORY_DIR    = Path("chat_histories")
ALLOWED_UPLOAD_EXTENSIONS = {".pdf", ".md", ".markdown", ".txt", ".docx", ".doc"}


# ── Tenant helpers ─────────────────────────────────────────────────────────────

def load_tenants() -> list:
    if not TENANTS_FILE.exists():
        return []
    try:
        return json.loads(TENANTS_FILE.read_text("utf-8"))
    except (json.JSONDecodeError, OSError):
        return []


def save_tenants(tenants: list) -> None:
    TENANTS_FILE.write_text(json.dumps(tenants, indent=2, ensure_ascii=False), encoding="utf-8")


def find_tenant(tenant_id: str) -> dict | None:
    return next((t for t in load_tenants() if t["id"] == tenant_id), None)


def tenant_stats(tenant: dict) -> dict:
    """Count chat sessions and data files for a tenant."""
    HISTORY_DIR.mkdir(exist_ok=True)
    sessions = list(HISTORY_DIR.glob(f"chat_history_*.json"))
    datasets = list(Path(DATA_PATH).glob("*.md")) if Path(DATA_PATH).exists() else []
    return {
        "session_count": len(sessions),
        "dataset_count": len(datasets),
    }


# ── Session helpers ────────────────────────────────────────────────────────────

def get_session_id():
    if "chat_session_id" not in session:
        session["chat_session_id"] = uuid4().hex
    return session["chat_session_id"]


def get_chat_history():
    return CHAT_SESSIONS.setdefault(get_session_id(), [])


def append_chat_turn(query, result):
    history = get_chat_history()
    history.append({"query": query, "result": result})
    del history[:-MAX_CHAT_TURNS]


def get_available_datasets():
    data_dir = Path(DATA_PATH)
    if not data_dir.exists():
        return []
    return sorted([f.name for f in data_dir.glob("*.md")])


def get_selected_dataset():
    return session.get("selected_dataset", None)


# ══════════════════════════════════════════════════════════════════════════════
# DASHBOARD ROUTES
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/")
def dashboard():
    tenants = load_tenants()
    stats   = {t["id"]: tenant_stats(t) for t in tenants}
    return render_template("dashboard.html", tenants=tenants, stats=stats)


# ── Tenant API ─────────────────────────────────────────────────────────────────

@app.route("/api/tenants", methods=["GET"])
def api_list_tenants():
    return jsonify(load_tenants())


@app.route("/api/tenants", methods=["POST"])
def api_create_tenant():
    data    = request.get_json(force=True) or {}
    name    = (data.get("name") or "").strip()
    website = (data.get("website") or "").strip()
    if not name:
        return jsonify({"error": "name is required"}), 400

    tenant = {
        "id":         uuid4().hex[:12],
        "name":       name,
        "website":    website,
        "api_key":    "ak_" + secrets.token_urlsafe(24),
        "created_at": __import__("datetime").datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ"),
        "notes":      data.get("notes", ""),
    }
    tenants = load_tenants()
    tenants.append(tenant)
    save_tenants(tenants)
    return jsonify(tenant), 201


@app.route("/api/tenants/<tenant_id>", methods=["DELETE"])
def api_delete_tenant(tenant_id):
    tenants = load_tenants()
    updated = [t for t in tenants if t["id"] != tenant_id]
    if len(updated) == len(tenants):
        return jsonify({"error": "not found"}), 404
    save_tenants(updated)
    return jsonify({"deleted": True})


@app.route("/api/tenants/<tenant_id>/regenerate-key", methods=["POST"])
def api_regenerate_key(tenant_id):
    tenants = load_tenants()
    for t in tenants:
        if t["id"] == tenant_id:
            t["api_key"] = "ak_" + secrets.token_urlsafe(24)
            save_tenants(tenants)
            return jsonify({"api_key": t["api_key"]})
    return jsonify({"error": "not found"}), 404


# ══════════════════════════════════════════════════════════════════════════════
# STUDIO ROUTES  (existing RAG test UI, moved to /studio)
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/studio", methods=["GET", "POST"])
def studio():
    result       = None
    query        = None
    ingest       = None
    chat_history = get_chat_history()
    if chat_history:
        result = chat_history[-1].get("result")

    if request.method == "POST":
        query    = request.form.get("query", "").strip()
        selected = get_selected_dataset()
        if query:
            result = query_data.query_rag_web(
                query,
                chat_history=chat_history,
                dataset_filter=selected,
            )
            append_chat_turn(query, result)

    return render_template(
        "studio.html",
        result=result,
        query=query,
        chat_history=chat_history,
        ingest=ingest,
        url="",
        model=MODEL_NAME,
        datasets=get_available_datasets(),
        selected_dataset=get_selected_dataset(),
    )


@app.route("/select-dataset", methods=["POST"])
def select_dataset():
    chosen    = request.form.get("dataset", "").strip()
    available = get_available_datasets()
    if chosen == "__all__" or chosen not in available:
        session.pop("selected_dataset", None)
    else:
        session["selected_dataset"] = chosen
    CHAT_SESSIONS[get_session_id()] = []
    return redirect(url_for("studio"))


@app.route("/clear-chat", methods=["POST"])
def clear_chat():
    CHAT_SESSIONS[get_session_id()] = []
    return redirect(url_for("studio"))


@app.route("/images/<path:filename>")
def images(filename):
    return send_from_directory("images", filename)


@app.route("/ingest-url", methods=["POST"])
def ingest_url():
    url    = request.form.get("url", "").strip()
    ingest = None
    try:
        validate_url(url)
        filename    = filename_from_url(url)
        saved_paths = scrape_full_website(url, filename)
        documents   = load_documents()
        chunks      = split_documents(documents)
        add_to_chroma(chunks)
        query_data._DB = None
        ingest = {
            "ok":      True,
            "message": f"Scraped {len(saved_paths)} page(s) from {url} and indexed them.",
            "path":    ", ".join(saved_paths[:3]) + ("..." if len(saved_paths) > 3 else ""),
        }
    except Exception as exc:
        ingest = {"ok": False, "message": f"Could not scrape/index that URL: {exc}", "path": None}

    return render_template(
        "studio.html",
        result=None, query=None,
        chat_history=get_chat_history(),
        ingest=ingest, url=url,
        model=MODEL_NAME,
        datasets=get_available_datasets(),
        selected_dataset=get_selected_dataset(),
    )


@app.route("/upload-file", methods=["POST"])
def upload_file():
    uploaded_file = request.files.get("document")
    ingest        = None
    try:
        if not uploaded_file or not uploaded_file.filename:
            raise ValueError("choose a PDF, Word, text, or Markdown file")
        original_name = secure_filename(uploaded_file.filename)
        extension     = os.path.splitext(original_name)[1].lower()
        if extension not in ALLOWED_UPLOAD_EXTENSIONS:
            raise ValueError("only PDF, Word, text, and Markdown files are supported")
        os.makedirs(DATA_PATH, exist_ok=True)
        saved_path = save_upload_as_markdown(uploaded_file, original_name)
        documents  = load_documents()
        chunks     = split_documents(documents)
        add_to_chroma(chunks)
        query_data._DB = None
        ingest = {
            "ok":      True,
            "message": f"Uploaded, converted to Markdown, and indexed {original_name}.",
            "path":    saved_path,
        }
    except Exception as exc:
        ingest = {"ok": False, "message": f"Could not upload/index that file: {exc}", "path": None}

    return render_template(
        "studio.html",
        result=None, query=None,
        chat_history=get_chat_history(),
        ingest=ingest, url="",
        model=MODEL_NAME,
        datasets=get_available_datasets(),
        selected_dataset=get_selected_dataset(),
    )


@app.route("/history/clear-all", methods=["POST"])
def clear_all_histories():
    HISTORY_DIR.mkdir(exist_ok=True)
    deleted = 0
    for f in HISTORY_DIR.glob("chat_history_*.json"):
        try:
            f.unlink()
            deleted += 1
        except OSError:
            continue
    return jsonify({"deleted": True, "deleted_count": deleted})


# ── File helpers ───────────────────────────────────────────────────────────────

def save_upload_as_markdown(uploaded_file, original_name: str) -> str:
    extension     = Path(original_name).suffix.lower()
    markdown_name = unique_markdown_filename(Path(original_name).stem)
    markdown_path = os.path.join(DATA_PATH, markdown_name)

    if extension in {".md", ".markdown"}:
        content = read_uploaded_text(uploaded_file)
    elif extension == ".txt":
        content = text_to_markdown(read_uploaded_text(uploaded_file), original_name)
    elif extension == ".pdf":
        content = pdf_to_markdown(uploaded_file, original_name)
    elif extension in {".docx", ".doc"}:
        content = docx_to_markdown(uploaded_file, original_name)
    else:
        raise ValueError("unsupported file type")

    with open(markdown_path, "w", encoding="utf-8") as f:
        f.write(content.strip() + "\n")
    return markdown_path


def unique_markdown_filename(base_name: str) -> str:
    safe_base = secure_filename(base_name) or "uploaded-document"
    candidate = f"{safe_base}.md"
    counter   = 2
    while os.path.exists(os.path.join(DATA_PATH, candidate)):
        candidate = f"{safe_base}-{counter}.md"
        counter  += 1
    return candidate


def read_uploaded_text(uploaded_file) -> str:
    raw = uploaded_file.read()
    for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def text_to_markdown(text: str, original_name: str) -> str:
    title = Path(original_name).stem.replace("_", " ").replace("-", " ").strip()
    return f"# {title or 'Uploaded document'}\n\n{text.strip()}"


def pdf_to_markdown(uploaded_file, original_name: str) -> str:
    uploaded_file.stream.seek(0)
    reader   = PdfReader(uploaded_file.stream)
    title    = Path(original_name).stem.replace("_", " ").replace("-", " ").strip()
    sections = [f"# {title or 'Uploaded PDF'}"]
    for i, page in enumerate(reader.pages, 1):
        text = (page.extract_text() or "").strip()
        if text:
            sections.append(f"## Page {i}\n\n{text}")
    if len(sections) == 1:
        raise ValueError("could not extract readable text from that PDF")
    return "\n\n".join(sections)


def docx_to_markdown(uploaded_file, original_name: str) -> str:
    if Path(original_name).suffix.lower() == ".doc":
        raise ValueError("legacy .doc files are not supported; please save it as .docx")
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("Word uploads need python-docx installed") from exc
    uploaded_file.stream.seek(0)
    doc   = Document(uploaded_file.stream)
    title = Path(original_name).stem.replace("_", " ").replace("-", " ").strip()
    lines = [f"# {title or 'Uploaded Word document'}"]
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                lines.append("| " + " | ".join(cells) + " |")
    if len(lines) == 1:
        raise ValueError("could not extract readable text from that Word document")
    return "\n\n".join(lines)


def validate_url(url: str):
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        raise ValueError("enter a full http:// or https:// URL")


def filename_from_url(url: str):
    parsed = urlparse(url)
    domain = parsed.netloc.replace("www.", "")
    path   = parsed.path.strip("/").replace("/", "-")
    base   = f"{domain}-{path}" if path else domain
    safe   = "".join(c if c.isalnum() or c in "-_." else "-" for c in base)
    return safe[:80]


if __name__ == "__main__":
    app.run(debug=True, port=5000)